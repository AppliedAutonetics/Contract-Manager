import os
import re
import json
import difflib
from datetime import datetime, date
from io import BytesIO

from flask import (Flask, render_template, redirect, url_for, flash, request,
                   jsonify, send_file, abort, session, make_response)
from markupsafe import Markup, escape
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect, CSRFError
from werkzeug.utils import secure_filename

from models import db, User, Client, ContractTemplate, Contract, ContractRevision, ContractFieldValue, AuditLog

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Word doc handling
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# mammoth: docx → HTML with formatting preserved
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

# xhtml2pdf: HTML → PDF
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False

# BeautifulSoup: HTML parsing for DOCX export
try:
    from bs4 import BeautifulSoup, NavigableString, Tag
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Mammoth style map: preserves heading levels and inline formatting from Word
MAMMOTH_STYLE_MAP = """
p[style-name='Heading 1'] => h1:fresh
p[style-name='Heading 2'] => h2:fresh
p[style-name='Heading 3'] => h3:fresh
p[style-name='Heading 4'] => h4:fresh
p[style-name='Heading 5'] => h5:fresh
p[style-name='Heading 6'] => h6:fresh
p[style-name='heading 1'] => h1:fresh
p[style-name='heading 2'] => h2:fresh
p[style-name='heading 3'] => h3:fresh
p[style-name='heading 4'] => h4:fresh
p[style-name='heading 5'] => h5:fresh
p[style-name='heading 6'] => h6:fresh
p[style-name='Title'] => h1.doc-title:fresh
p[style-name='Subtitle'] => p.doc-subtitle:fresh
p[style-name='Quote'] => blockquote:fresh
p[style-name='Intense Quote'] => blockquote:fresh
p[style-name='Body Text'] => p:fresh
p[style-name='Body Text 2'] => p:fresh
p[style-name='Body Text 3'] => p:fresh
p[style-name='Body Text Indent'] => p.list-paragraph:fresh
p[style-name='Body Text Indent 2'] => p.list-paragraph:fresh
p[style-name='List Paragraph'] => p.list-paragraph:fresh
p[style-name='List Bullet'] => p:fresh
p[style-name='List Bullet 2'] => p:fresh
p[style-name='List Bullet 3'] => p:fresh
p[style-name='List Number'] => p:fresh
p[style-name='List Number 2'] => p:fresh
p[style-name='List Number 3'] => p:fresh
p[style-name='Caption'] => p.caption:fresh
p[style-name='TOC 1'] => p:fresh
p[style-name='TOC 2'] => p.list-paragraph:fresh
p[style-name='TOC 3'] => p.list-paragraph:fresh
p[style-name='No Spacing'] => p:fresh
p[style-name='Normal (Web)'] => p:fresh
p[style-name='Text Body'] => p:fresh
p[style-name='First Paragraph'] => p:fresh
p[style-name='Normal Indent'] => p.list-paragraph:fresh
r[style-name='Strong'] => strong
r[style-name='Emphasis'] => em
r[style-name='Intense Emphasis'] => em
r[style-name='Book Title'] => strong
r[style-name='Subtle Reference'] => em
r[style-name='Intense Reference'] => strong
b => strong
i => em
u => u
strike => s
"""

app = Flask(__name__)

# Respect X-Forwarded-* headers from nginx/reverse-proxy so Flask sees the
# real scheme and host.  x_for=1 trusts one proxy hop (nginx → gunicorn).
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# Configuration
# Guard against SESSION_SECRET being set to an empty string in the environment.
_secret = os.environ.get('SESSION_SECRET') or 'dev-secret-key-change-in-prod'
app.config['SECRET_KEY'] = _secret
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'postgresql://localhost/contracts')
if app.config['SQLALCHEMY_DATABASE_URI'].startswith('postgres://'):
    app.config['SQLALCHEMY_DATABASE_URI'] = app.config['SQLALCHEMY_DATABASE_URI'].replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload
# Disable CSRF token expiry — the default 1-hour limit causes spurious
# failures when a page stays open or a reverse proxy delays the request.
app.config['WTF_CSRF_TIME_LIMIT'] = None
# Ensure session cookies work over plain HTTP on local/production servers.
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize extensions
db.init_app(app)
csrf = CSRFProtect(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'info'


@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    """Show a friendly message instead of a raw 400 and send the user back."""
    flash('Your session expired or the form token was invalid — please try again.', 'error')
    return redirect(request.referrer or url_for('dashboard')), 302


@app.template_filter('render_content')
def render_content_filter(content):
    """Render contract/template content safely.

    HTML content (from mammoth .docx conversion) is returned as Markup so
    Jinja2 doesn't escape it.  Plain-text content is wrapped in a <pre> so
    line breaks and spacing are preserved.
    """
    if not content:
        return Markup('')
    if is_html_content(content):
        return Markup(content)
    return Markup('<pre style="white-space:pre-wrap;font-family:inherit">' +
                  str(escape(content)) + '</pre>')


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ─── Helpers ──────────────────────────────────────────────────────────────────

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


_PLUS_PLACEHOLDER_RE = re.compile(r'\+([A-Za-z_][A-Za-z0-9_]*)\+')


def _normalize_plus_placeholders(text):
    """Convert +FIELD_NAME+ / +field_name+ → {{FIELD_NAME}} (uppercased).

    Many Word templates use +NAME+ notation instead of {{NAME}}.  This
    normalises them to our canonical format so field extraction and
    substitution work correctly for both styles.
    """
    return _PLUS_PLACEHOLDER_RE.sub(lambda m: '{{' + m.group(1).upper() + '}}', text)


def fix_split_placeholders(html):
    """Strip HTML tags that Word/mammoth inserted inside {{...}} placeholders
    and normalise variant placeholder formats.

    Word sometimes stores a placeholder like {{CLIENT_NAME}} as multiple runs
    with mixed formatting, so mammoth wraps each run separately producing
    things like {{<strong>CLIENT</strong>_NAME}}.  This step strips inline tags
    found between {{ and }} so the markers are clean.

    It also:
    - Normalises HTML-entity-encoded braces  (&#123;&#123; → {{)
    - Normalises +FIELD_NAME+ notation       (→ {{FIELD_NAME}})
    - Stops at block-level tag boundaries so it never merges across cells
    """
    if not html:
        return html

    # Normalise HTML-encoded {{ and }}
    html = html.replace('&#123;&#123;', '{{').replace('&#125;&#125;', '}}')
    html = html.replace('&lbrace;&lbrace;', '{{').replace('&rbrace;&rbrace;', '}}')
    html = html.replace('&#x7B;&#x7B;', '{{').replace('&#x7D;&#x7D;', '}}')

    # Normalise +FIELD_NAME+ → {{FIELD_NAME}} before any other processing
    html = _normalize_plus_placeholders(html)

    if '{{' not in html:
        return html

    _BLOCK_STOP = re.compile(
        r'</?(?:p|div|td|th|tr|li|h[1-6]|table|blockquote|section|article)\b',
        re.IGNORECASE,
    )

    def _replacer(m):
        content = m.group(1)
        # Don't merge if there's a block-level boundary inside the match
        if _BLOCK_STOP.search(content):
            return m.group(0)
        cleaned = re.sub(r'<[^>]+>', '', content).strip()
        return '{{' + cleaned + '}}'

    return re.sub(
        r'\{\{((?:[^{}]|<[^>]+>)*?)\}\}',
        _replacer,
        html,
        flags=re.DOTALL,
    )


def extract_template_fields(content):
    """Find all {{FIELD_NAME}} markers in template content (HTML or plain text).

    Also detects +FIELD_NAME+ notation by normalising first, so both formats
    are transparently supported.
    """
    # Normalise +FIELD+ notation in plain-text content before searching
    if not is_html_content(content):
        content = _normalize_plus_placeholders(content)

    pattern = r'\{\{([A-Z_][A-Z0-9_]*)\}\}'
    if is_html_content(content):
        # Clean split placeholders first, then also search the plain-text
        # version to catch any that are still split across separate elements.
        cleaned = fix_split_placeholders(content)
        import html as _html_mod
        plain = _html_mod.unescape(re.sub(r'<[^>]+>', '', cleaned))
        fields = list(dict.fromkeys(
            re.findall(pattern, cleaned) + re.findall(pattern, plain)
        ))
        return fields
    return list(dict.fromkeys(re.findall(pattern, content)))


def apply_field_values(content, field_values):
    """Replace {{FIELD_NAME}} markers (and +FIELD_NAME+ variants) with actual values.

    For HTML content the split-placeholder cleaner runs first so that markers
    like {{<strong>CLIENT_NAME</strong>}} are normalised before substitution.
    Unfilled fields (no value provided) are left as the original {{FIELD_NAME}}
    placeholder so the document clearly shows what is still missing.
    """
    # Normalise +FIELD+ notation in plain-text content
    if not is_html_content(content):
        content = _normalize_plus_placeholders(content)
    else:
        content = fix_split_placeholders(content)   # also normalises +FIELD+

    for name, value in field_values.items():
        placeholder = '{{' + name + '}}'
        content = content.replace(placeholder, value if value else placeholder)
    return content


def generate_contract_number(offset=0):
    """Return the next unused CTR-YYYY-NNNN contract number.

    Uses MAX of existing sequence numbers rather than COUNT so that
    deletions or gaps never produce a duplicate.  Pass offset>0 to skip
    ahead when the caller detects a collision and needs to retry.
    """
    year = datetime.utcnow().year
    prefix = f'CTR-{year}-'
    rows = db.session.execute(
        db.text("SELECT contract_number FROM contracts WHERE contract_number LIKE :p"),
        {'p': prefix + '%'}
    ).fetchall()
    max_seq = 0
    for (num,) in rows:
        try:
            seq = int(num[len(prefix):])
            if seq > max_seq:
                max_seq = seq
        except (ValueError, IndexError):
            pass
    return f'{prefix}{max_seq + 1 + offset:04d}'


def log_action(action, resource_type=None, resource_id=None, contract_id=None, details=None):
    log = AuditLog(
        user_id=current_user.id if current_user.is_authenticated else None,
        action=action,
        resource_type=resource_type,
        resource_id=resource_id,
        contract_id=contract_id,
        details=details,
        ip_address=request.remote_addr,
        created_at=datetime.utcnow()
    )
    db.session.add(log)


def is_html_content(content):
    """Return True when content is HTML (produced by mammoth from a .docx)."""
    return bool(content and content.strip().startswith('<'))


def _docx_to_html(file_path):
    """Convert a .docx to HTML with full inline CSS (colors, sizes, alignment).

    Uses python-docx directly to read the Word XML so that run-level formatting
    (text color, font size, bold/italic/underline) and paragraph-level formatting
    (alignment, spacing, indentation) are preserved as inline CSS rather than
    being discarded.  Returns an HTML string.
    """
    import html as _hl

    try:
        from docx import Document as _Doc
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        return None  # caller will fall back to mammoth

    try:
        doc = _Doc(file_path)
    except Exception as e:
        return f'<p>[Could not open Word document: {_hl.escape(str(e))}]</p>'

    _H_STYLES = {
        'heading 1': 'h1', 'heading 2': 'h2', 'heading 3': 'h3',
        'heading 4': 'h4', 'heading 5': 'h5', 'heading 6': 'h6',
        'title': 'h1', 'subtitle': 'h2',
    }

    def _run_html(run):
        text = _hl.escape(run.text or '')
        if not text:
            return ''

        span_styles = []

        # Text colour
        try:
            if run.font.color and run.font.color.type is not None:
                rgb = str(run.font.color.rgb)   # e.g. '1A2742'
                span_styles.append(f'color:#{rgb.lower()}')
        except Exception:
            pass

        # Font size
        try:
            if run.font.size:
                pts = run.font.size.pt
                if pts:
                    span_styles.append(f'font-size:{pts:.1f}pt')
        except Exception:
            pass

        # Font family (skip common defaults to keep HTML clean)
        try:
            fname = run.font.name
            if fname and fname.strip() not in ('', 'Calibri', 'Arial', 'Times New Roman',
                                               'Cambria', 'Calibri Light'):
                span_styles.append(f"font-family:'{_hl.escape(fname)}',sans-serif")
        except Exception:
            pass

        # Semantic inline tags
        if run.bold:      text = f'<strong>{text}</strong>'
        if run.italic:    text = f'<em>{text}</em>'
        if run.underline: text = f'<u>{text}</u>'
        try:
            if run.font.strike: text = f'<s>{text}</s>'
        except Exception:
            pass
        try:
            if run.font.superscript: text = f'<sup>{text}</sup>'
            elif run.font.subscript: text = f'<sub>{text}</sub>'
        except Exception:
            pass

        if span_styles:
            text = f'<span style="{";".join(span_styles)}">{text}</span>'

        return text

    def _para_styles(para):
        """Return (tag, style_attr, list_type_or_None)."""
        sname = (para.style.name or 'Normal').strip().lower()
        tag = _H_STYLES.get(sname, 'p')

        pstyles = []
        pf = para.paragraph_format

        # Alignment
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WDA
            _amap = {_WDA.CENTER: 'center', _WDA.RIGHT: 'right', _WDA.JUSTIFY: 'justify'}
            a = _amap.get(pf.alignment)
            if a:
                pstyles.append(f'text-align:{a}')
        except Exception:
            pass

        # Vertical spacing
        try:
            if pf.space_before and pf.space_before.pt:
                pstyles.append(f'margin-top:{pf.space_before.pt:.0f}pt')
            if pf.space_after and pf.space_after.pt:
                pstyles.append(f'margin-bottom:{pf.space_after.pt:.0f}pt')
        except Exception:
            pass

        # Indentation
        try:
            if pf.left_indent and pf.left_indent.inches and pf.left_indent.inches > 0:
                pstyles.append(f'padding-left:{pf.left_indent.inches:.3f}in')
        except Exception:
            pass

        style_attr = f' style="{";".join(pstyles)}"' if pstyles else ''

        # List detection
        list_type = None
        try:
            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            numPr = para._element.find(f'{{{W}}}pPr/{{{W}}}numPr')
            if numPr is not None:
                # Try to determine bullet vs numbered from abstract num
                ilvl = numPr.find(f'{{{W}}}ilvl')
                numId = numPr.find(f'{{{W}}}numId')
                if numId is not None and numId.get(f'{{{W}}}val', '0') != '0':
                    # Heuristic: styles with "number" are numbered lists
                    list_type = 'number' if 'number' in sname else 'bullet'
        except Exception:
            pass
        if list_type is None:
            if 'list bullet' in sname:   list_type = 'bullet'
            elif 'list number' in sname: list_type = 'number'

        return tag, style_attr, list_type

    def _para_html(para, force_tag=None, style_attr=None):
        if style_attr is None:
            tag, style_attr, _ = _para_styles(para)
        else:
            tag = force_tag or 'p'
        inner = ''.join(_run_html(r) for r in para.runs)
        if not inner.strip():
            return f'<{tag}{style_attr}><br></{tag}>'
        return f'<{tag}{style_attr}>{inner}</{tag}>'

    def _table_html(table):
        # Detect header row: row 0 is a header if all its first-paragraph runs are bold
        def _row_is_header(row):
            try:
                for cell in row.cells:
                    if not cell.paragraphs:
                        return False
                    runs = [r for r in cell.paragraphs[0].runs if r.text.strip()]
                    if runs and not all(r.bold for r in runs):
                        return False
                return bool(row.cells)
            except Exception:
                return False

        rows = []
        header_detected = _row_is_header(table.rows[0]) if table.rows else False
        for ri, row in enumerate(table.rows):
            cells_html = []
            seen_ids = set()
            is_header_row = (ri == 0 and header_detected)
            ctag = 'th' if is_header_row else 'td'
            for cell in row.cells:
                cid = id(cell._tc)
                if cid in seen_ids:
                    continue        # skip merged duplicates
                seen_ids.add(cid)
                # Collect non-empty paragraph HTML; skip trailing blank paras
                cell_parts = []
                for p in cell.paragraphs:
                    inner = ''.join(_run_html(r) for r in p.runs)
                    if inner.strip():
                        tag, sattr, _ = _para_styles(p)
                        cell_parts.append(f'<{tag}{sattr}>{inner}</{tag}>')
                if not cell_parts:
                    cell_parts = ['<p>&nbsp;</p>']
                cells_html.append(f'<{ctag}>{"".join(cell_parts)}</{ctag}>')
            rows.append(f'<tr>{"".join(cells_html)}</tr>')
        return f'<table>{"".join(rows)}</table>'

    # Walk the body in document order, grouping consecutive list paragraphs
    parts = []
    pending_list = []       # list of (list_type, para, style_attr) tuples
    pending_list_type = None

    def flush_list():
        nonlocal pending_list, pending_list_type
        if not pending_list:
            return
        ltag = 'ul' if pending_list_type == 'bullet' else 'ol'
        items = ['<li>' + ''.join(_run_html(r) for r in lp.runs) + '</li>'
                 for (_, lp, _) in pending_list]
        parts.append(f'<{ltag}>{"".join(items)}</{ltag}>\n')
        pending_list = []
        pending_list_type = None

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for child in doc.element.body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if local == 'p':
            para = _Para(child, doc)
            tag, style_attr, list_type = _para_styles(para)

            if list_type:
                if list_type != pending_list_type and pending_list:
                    flush_list()
                pending_list_type = list_type
                pending_list.append((list_type, para, style_attr))
            else:
                flush_list()
                parts.append(_para_html(para, tag, style_attr) + '\n')

        elif local == 'tbl':
            flush_list()
            try:
                table = _Tbl(child, doc)
                parts.append(_table_html(table) + '\n')
            except Exception as te:
                parts.append(f'<p>[Table render error: {_hl.escape(str(te))}]</p>\n')

    flush_list()
    html = ''.join(parts)
    # Fix any placeholders Word split across runs
    return fix_split_placeholders(html)


def extract_text_from_file(file_path, filename):
    """Extract content from an uploaded file.

    For .docx files the python-docx converter is used first because it
    produces HTML with full inline CSS (colours, font sizes, alignment).
    Mammoth is kept as a fallback for files that python-docx can't open.
    The returned string is an HTML fragment; plain text for .txt files.
    """
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()

        elif ext in ('docx', 'doc'):
            # ── Primary: python-docx with full formatting ────────────────────
            if DOCX_AVAILABLE:
                html = _docx_to_html(file_path)
                if html and not html.startswith('<p>['):   # not an error result
                    return html

            # ── Fallback: mammoth (structural HTML, no colours) ──────────────
            if MAMMOTH_AVAILABLE:
                with open(file_path, 'rb') as f:
                    result = mammoth.convert_to_html(
                        f,
                        style_map=MAMMOTH_STYLE_MAP,
                        convert_image=mammoth.images.inline(
                            lambda image: {'src': ''}
                        ),
                    )
                html = fix_split_placeholders(result.value)
                return html

            # ── Last resort: plain text via python-docx ──────────────────────
            if DOCX_AVAILABLE:
                doc = DocxDocument(file_path)
                return '\n'.join(para.text for para in doc.paragraphs)

            return '[Word document: install python-docx or mammoth to extract text]'

        else:
            return f'[File: {filename} — content extraction not supported for this format]'

    except Exception as e:
        return f'[Error reading file: {str(e)}]'


# ─── PDF Generation ───────────────────────────────────────────────────────────

def generate_pdf(contract, revision=None):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'ContractTitle',
        parent=styles['Title'],
        fontSize=20,
        spaceAfter=6,
        textColor=colors.HexColor('#1a2742'),
        alignment=TA_CENTER
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=4,
        textColor=colors.HexColor('#64748b'),
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=13,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor('#1a2742'),
        borderPad=4
    )
    body_style = ParagraphStyle(
        'ContractBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=16,
        spaceAfter=8,
        textColor=colors.HexColor('#374151')
    )
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#374151')
    )

    # Header
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(contract.title, title_style))
    story.append(Paragraph(f'Contract #{contract.contract_number}', subtitle_style))
    story.append(Spacer(1, 0.1 * inch))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#1a2742')))
    story.append(Spacer(1, 0.2 * inch))

    # Contract details table
    status_map = {'draft': 'Draft', 'in_review': 'In Review', 'approved': 'Approved', 'finalized': 'Finalized', 'expired': 'Expired'}
    details_data = [
        ['Client:', contract.client.name, 'Status:', status_map.get(contract.status, contract.status)],
        ['Created:', contract.created_at.strftime('%B %d, %Y'), 'Creator:', contract.creator.full_name],
    ]
    if contract.start_date:
        details_data.append(['Start Date:', contract.start_date.strftime('%B %d, %Y'), 'End Date:', contract.end_date.strftime('%B %d, %Y') if contract.end_date else 'N/A'])
    if contract.value:
        details_data.append(['Contract Value:', f'${float(contract.value):,.2f}', '', ''])

    details_table = Table(details_data, colWidths=[1.3 * inch, 2.5 * inch, 1.3 * inch, 2.5 * inch])
    details_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#374151')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1a2742')),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#1a2742')),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#f8fafc'), colors.white]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(details_table)
    story.append(Spacer(1, 0.3 * inch))

    # Contract content
    content = ''
    if revision:
        content = revision.content
        story.append(Paragraph(f'Version {revision.version_number}' + (' (FINAL)' if revision.is_finalized else ''), subtitle_style))
    elif contract.latest_revision:
        content = contract.latest_revision.content

    if content:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
        story.append(Spacer(1, 0.15 * inch))

        if is_html_content(content):
            # HTML content: delegate to xhtml2pdf which handles formatting
            doc.build(story)
            buffer.seek(0)
            return _generate_pdf_html(contract, content, revision)

        for para_text in content.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                story.append(Spacer(1, 0.1 * inch))
                continue

            # Detect section headings (ALL CAPS or ends with colon)
            lines = para_text.split('\n')
            first_line = lines[0].strip()
            if first_line.isupper() and len(first_line) < 80:
                story.append(Paragraph(first_line, heading_style))
                if len(lines) > 1:
                    rest = ' '.join(lines[1:]).strip()
                    if rest:
                        story.append(Paragraph(rest, body_style))
            else:
                # Normal paragraph - replace newlines with spaces
                text = para_text.replace('\n', ' ')
                story.append(Paragraph(text, body_style))

    # Footer info
    story.append(Spacer(1, 0.4 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 0.1 * inch))
    gen_time = datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')
    story.append(Paragraph(f'Generated on {gen_time}', subtitle_style))
    if contract.finalized_at:
        story.append(Paragraph(f'Finalized on {contract.finalized_at.strftime("%B %d, %Y")}', subtitle_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _generate_pdf_html(contract, html_body, revision=None):
    """Generate a PDF from the contract body only — no administrative metadata."""
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  @page {{ margin: 1in; }}
  body {{ font-family: Arial, sans-serif; font-size: 10pt; color: #374151; line-height: 1.6; }}

  /* Headings */
  h1 {{ font-size: 14pt; font-weight: bold; color: #1a2742; margin: 12pt 0 5pt 0; }}
  h2 {{ font-size: 12pt; font-weight: bold; color: #1a2742; margin: 10pt 0 4pt 0; }}
  h3 {{ font-size: 11pt; font-weight: bold; color: #374151; margin: 8pt 0 3pt 0; }}
  h4, h5, h6 {{ font-size: 10pt; font-weight: bold; margin: 6pt 0 2pt 0; }}

  /* Paragraphs & special styles */
  p {{ margin: 3pt 0 5pt 0; }}
  .doc-title {{ font-size: 16pt; font-weight: bold; text-align: center; color: #1a2742; margin: 0 0 4pt 0; }}
  .doc-subtitle {{ font-size: 11pt; text-align: center; color: #6b7280; margin: 2pt 0 10pt 0; }}
  .list-paragraph {{ margin-left: 18pt; }}
  .caption {{ font-size: 9pt; color: #6b7280; font-style: italic; text-align: center; }}

  /* Inline formatting */
  strong, b {{ font-weight: bold; }}
  em, i {{ font-style: italic; }}
  u {{ text-decoration: underline; }}
  s, strike, del {{ text-decoration: line-through; }}
  sub {{ vertical-align: sub; font-size: 8pt; }}
  sup {{ vertical-align: super; font-size: 8pt; }}
  a {{ color: #1d4ed8; text-decoration: underline; }}
  code, pre {{ font-family: Courier, monospace; font-size: 9pt; }}

  /* Blockquote */
  blockquote {{ margin: 6pt 0 6pt 20pt; padding-left: 8pt;
                border-left: 2pt solid #d1d5db; color: #6b7280; }}

  /* Tables */
  table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 9.5pt; }}
  td, th {{ border: 0.5pt solid #374151; padding: 4pt 7pt;
            vertical-align: top; word-break: break-word; }}
  th {{ font-weight: bold; background: #f1f5f9; color: #1a2742; }}

  /* Lists */
  ul, ol {{ margin: 3pt 0 6pt 0; padding-left: 18pt; }}
  li {{ margin-bottom: 2pt; line-height: 1.5; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    buf = BytesIO()
    if XHTML2PDF_AVAILABLE:
        pisa.CreatePDF(html, dest=buf)
    else:
        # Fallback: plain-text PDF via ReportLab (formatting lost)
        import html as html_lib
        plain = re.sub(r'<[^>]+>', ' ', html_body)
        plain = html_lib.unescape(plain)
        buf = BytesIO()
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        fb_doc = SimpleDocTemplate(buf, pagesize=letter,
                                   rightMargin=inch, leftMargin=inch,
                                   topMargin=inch, bottomMargin=inch)
        st = getSampleStyleSheet()
        story = []
        for para in plain.split('\n\n'):
            para = para.strip()
            if para:
                story.append(Paragraph(para.replace('\n', ' '), st['Normal']))
                story.append(Spacer(1, 8))
        fb_doc.build(story)

    buf.seek(0)
    return buf


def _get_para_align(element):
    """Read text-align from an element's style attribute and return WD_ALIGN_PARAGRAPH."""
    style = element.get('style', '')
    classes = element.get('class', [])
    if isinstance(classes, str):
        classes = classes.split()
    # Class-implied alignment
    if any(c in classes for c in ('doc-title', 'doc-subtitle', 'caption')):
        return WD_ALIGN_PARAGRAPH.CENTER
    # Inline style
    for part in style.replace(' ', '').split(';'):
        if part.startswith('text-align:'):
            val = part.split(':', 1)[1].lower()
            if val == 'center':  return WD_ALIGN_PARAGRAPH.CENTER
            if val == 'right':   return WD_ALIGN_PARAGRAPH.RIGHT
            if val == 'justify': return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None


def _parse_css_props(style_str):
    """Parse a CSS inline style string → dict of {prop: value}."""
    props = {}
    for decl in (style_str or '').split(';'):
        decl = decl.strip()
        if ':' in decl:
            k, _, v = decl.partition(':')
            props[k.strip().lower()] = v.strip()
    return props


def _css_color_to_rgb(value):
    """Convert a CSS colour value to an (r, g, b) int tuple, or None."""
    from docx.shared import RGBColor
    v = (value or '').strip().lower()
    if v.startswith('#'):
        h = v[1:]
        if len(h) == 3:
            h = h[0] * 2 + h[1] * 2 + h[2] * 2
        if len(h) == 6:
            try:
                return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
            except ValueError:
                pass
    elif v.startswith('rgb('):
        try:
            parts = v[4:v.index(')')].split(',')
            return tuple(int(p.strip()) for p in parts[:3])
        except Exception:
            pass
    _NAMED = {
        'black': (0, 0, 0), 'white': (255, 255, 255), 'red': (220, 38, 38),
        'blue': (37, 99, 235), 'green': (22, 163, 74), 'gray': (107, 114, 128),
        'grey': (107, 114, 128), 'navy': (26, 39, 66), 'orange': (234, 88, 12),
        'purple': (124, 58, 237), 'yellow': (202, 138, 4),
    }
    return _NAMED.get(v)


def _add_inline_to_para(para, node, bold=False, italic=False,
                         underline=False, strike=False,
                         color=None, font_size_pt=None):
    """Recursively add inline HTML into a docx paragraph, applying all formatting.

    Formatting flags (bold, italic, etc.) and CSS-derived properties (colour,
    font-size) accumulate as we descend so that nested markup like
    <strong><em><span style="color:#f00">text</span></em></strong> is handled
    correctly at every level.
    """
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = para.add_run(text)
            if bold:         run.bold        = True
            if italic:       run.italic      = True
            if underline:    run.underline   = True
            if strike:       run.font.strike = True
            if color:
                try:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(*color)
                except Exception:
                    pass
            if font_size_pt:
                try:
                    from docx.shared import Pt
                    run.font.size = Pt(font_size_pt)
                except Exception:
                    pass
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()

    if tag == 'br':
        para.add_run('\n')
        return

    # Accumulate semantic formatting flags
    _bold      = bold      or tag in ('strong', 'b')
    _italic    = italic    or tag in ('em', 'i')
    _underline = underline or tag == 'u'
    _strike    = strike    or tag in ('s', 'strike', 'del')
    _color     = color
    _font_size = font_size_pt

    # Parse inline CSS from any element (most commonly <span style="...">)
    style_str = node.get('style', '')
    if style_str:
        props = _parse_css_props(style_str)
        if 'color' in props:
            c = _css_color_to_rgb(props['color'])
            if c:
                _color = c
        if 'font-size' in props:
            fs = props['font-size']
            try:
                if fs.endswith('pt'):
                    _font_size = float(fs[:-2])
                elif fs.endswith('px'):
                    _font_size = round(float(fs[:-2]) * 0.75, 1)
            except ValueError:
                pass

    for child in node.children:
        _add_inline_to_para(para, child, _bold, _italic, _underline, _strike,
                            _color, _font_size)


def _html_to_docx_body(doc, html_content):
    """Parse an HTML fragment and append its content to a python-docx Document."""
    if not BS4_AVAILABLE:
        import html as _hl
        plain = _hl.unescape(re.sub(r'<[^>]+>', ' ', html_content))
        for para in plain.split('\n'):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        return
    soup = BeautifulSoup(html_content, 'lxml')
    body = soup.find('body') or soup
    _add_block_children(doc, body, left_indent=None)


def _add_block_children(doc, parent, left_indent=None):
    """Walk direct children of a block element, dispatching each to _add_block_element."""
    for element in parent.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = doc.add_paragraph(text)
                if left_indent is not None:
                    p.paragraph_format.left_indent = left_indent
            continue
        if isinstance(element, Tag):
            _add_block_element(doc, element, left_indent=left_indent)


def _add_block_element(doc, element, left_indent=None):
    """Convert one block-level HTML element into python-docx content."""
    tag = element.name.lower()
    classes = element.get('class', [])
    if isinstance(classes, str):
        classes = classes.split()

    # ── Headings ────────────────────────────────────────────────────────────
    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        para = doc.add_heading('', level=level)
        for child in element.children:
            _add_inline_to_para(para, child)
        align = _get_para_align(element)
        if align is not None:
            para.alignment = align
        if left_indent is not None:
            para.paragraph_format.left_indent = left_indent

    # ── Paragraph ───────────────────────────────────────────────────────────
    elif tag == 'p':
        para = doc.add_paragraph()
        for child in element.children:
            _add_inline_to_para(para, child)
        align = _get_para_align(element)
        if align is not None:
            para.alignment = align
        if 'list-paragraph' in classes:
            para.paragraph_format.left_indent = Inches(0.5)
        elif left_indent is not None:
            para.paragraph_format.left_indent = left_indent

    # ── Lists ────────────────────────────────────────────────────────────────
    elif tag in ('ul', 'ol'):
        style = 'List Bullet' if tag == 'ul' else 'List Number'
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style=style)
            for child in li.children:
                _add_inline_to_para(para, child)
            if left_indent is not None:
                para.paragraph_format.left_indent = left_indent

    # ── Table ────────────────────────────────────────────────────────────────
    elif tag == 'table':
        rows_html = element.find_all('tr')
        if not rows_html:
            return
        max_cols = max((len(r.find_all(['td', 'th'])) for r in rows_html), default=0)
        if max_cols == 0:
            return
        tbl = doc.add_table(rows=len(rows_html), cols=max_cols)
        tbl.style = 'Table Grid'
        for ri, row_el in enumerate(rows_html):
            for ci, cell_el in enumerate(row_el.find_all(['td', 'th'])):
                if ci >= max_cols:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = ''
                is_header = cell_el.name == 'th'

                # Get direct <p>/<hN> block children (each becomes its own paragraph)
                block_children = [c for c in cell_el.children
                                  if isinstance(c, Tag) and c.name in
                                  ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6')]

                if block_children:
                    para = cell.paragraphs[0]
                    for pi, blk in enumerate(block_children):
                        if pi > 0:
                            para = cell.add_paragraph()
                        for child in blk.children:
                            _add_inline_to_para(para, child, bold=is_header)
                        c_align = _get_para_align(blk)
                        if c_align is not None:
                            para.alignment = c_align
                        if is_header:
                            try:
                                from docx.shared import RGBColor as _RGB
                                for run in para.runs:
                                    if not run.font.color.type:
                                        run.font.color.rgb = _RGB(0x1a, 0x27, 0x42)
                            except Exception:
                                pass
                else:
                    # No block wrappers — treat all children as inline content
                    para = cell.paragraphs[0]
                    for child in cell_el.children:
                        _add_inline_to_para(para, child, bold=is_header)
                    c_align = _get_para_align(cell_el)
                    if c_align is not None:
                        para.alignment = c_align

    # ── Blockquote — indented container ──────────────────────────────────────
    elif tag == 'blockquote':
        _add_block_children(doc, element, left_indent=Inches(0.5))

    # ── Generic containers ───────────────────────────────────────────────────
    elif tag in ('div', 'section', 'article', 'main', 'aside', 'figure'):
        _add_block_children(doc, element, left_indent=left_indent)

    # ── Horizontal rule ──────────────────────────────────────────────────────
    elif tag == 'hr':
        doc.add_paragraph('─' * 60)

    # ── Line break as spacing paragraph ──────────────────────────────────────
    elif tag == 'br':
        doc.add_paragraph()

    # Skip: script, style, head, noscript, etc.


def generate_docx(contract, revision=None):
    """Generate a Word (.docx) containing only the contract body.

    Document-level styles are set to match the PDF output (Arial 10 pt,
    colour #374151 for body; matching size/colour/weight for headings) so
    that DOCX and PDF exports look identical to the reader.
    """
    if not DOCX_AVAILABLE:
        return None

    doc = DocxDocument()

    # ── Standard 1-inch page margins ────────────────────────────────────────
    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # ── Base style — matches PDF: Arial 10 pt, colour #374151 ───────────────
    try:
        normal = doc.styles['Normal']
        normal.font.name      = 'Arial'
        normal.font.size      = Pt(10)
        normal.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    except Exception:
        pass

    # ── Heading styles — match PDF CSS ───────────────────────────────────────
    _HEADING_DEFS = {
        1: (Pt(14), RGBColor(0x1a, 0x27, 0x42)),
        2: (Pt(12), RGBColor(0x1a, 0x27, 0x42)),
        3: (Pt(11), RGBColor(0x37, 0x41, 0x51)),
        4: (Pt(10), RGBColor(0x37, 0x41, 0x51)),
        5: (Pt(10), RGBColor(0x37, 0x41, 0x51)),
        6: (Pt(10), RGBColor(0x37, 0x41, 0x51)),
    }
    for level, (size, color) in _HEADING_DEFS.items():
        try:
            h = doc.styles[f'Heading {level}']
            h.font.name      = 'Arial'
            h.font.size      = size
            h.font.color.rgb = color
            h.font.bold      = True
        except Exception:
            pass

    # ── List styles — consistent font ────────────────────────────────────────
    for lst_style in ('List Bullet', 'List Number', 'List Paragraph'):
        try:
            s = doc.styles[lst_style]
            s.font.name = 'Arial'
            s.font.size = Pt(10)
        except Exception:
            pass

    # ── Contract content only ────────────────────────────────────────────────
    content = ''
    if revision:
        content = revision.content
    elif contract.latest_revision:
        content = contract.latest_revision.content

    if content:
        if is_html_content(content):
            _html_to_docx_body(doc, content)
        else:
            for para_text in content.split('\n\n'):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Auth Routes ──────────────────────────────────────────────────────────────

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        remember = request.form.get('remember') == 'on'

        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            user.last_login = datetime.utcnow()
            db.session.commit()
            login_user(user, remember=remember)
            next_page = request.args.get('next')
            if next_page:
                from urllib.parse import urlparse
                parsed = urlparse(next_page)
                if parsed.scheme or parsed.netloc:  # absolute URL — reject it
                    next_page = None
            flash(f'Welcome back, {user.full_name}!', 'success')
            return redirect(next_page or url_for('dashboard'))
        else:
            flash('Invalid email or password.', 'error')

    return render_template('auth/login.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        full_name = request.form.get('full_name', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if not all([full_name, email, password]):
            flash('All fields are required.', 'error')
        elif password != confirm:
            flash('Passwords do not match.', 'error')
        elif len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
        elif User.query.filter_by(email=email).first():
            flash('An account with this email already exists.', 'error')
        else:
            user = User(full_name=full_name, email=email)
            user.set_password(password)
            db.session.add(user)
            db.session.commit()
            login_user(user)
            flash(f'Account created! Welcome, {full_name}.', 'success')
            return redirect(url_for('dashboard'))

    return render_template('auth/register.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


# ─── Dashboard ────────────────────────────────────────────────────────────────

@app.route('/')
@login_required
def dashboard():
    total_contracts = Contract.query.filter_by(created_by=current_user.id).count()
    by_status = {
        'draft':              Contract.query.filter_by(created_by=current_user.id, status='draft').count(),
        'in_review':          Contract.query.filter_by(created_by=current_user.id, status='in_review').count(),
        'approved':           Contract.query.filter_by(created_by=current_user.id, status='approved').count(),
        'finalized':          Contract.query.filter_by(created_by=current_user.id, status='finalized').count(),
        'partially_executed': Contract.query.filter_by(created_by=current_user.id, status='partially_executed').count(),
        'fully_executed':     Contract.query.filter_by(created_by=current_user.id, status='fully_executed').count(),
    }
    recent_contracts = Contract.query.filter_by(created_by=current_user.id).order_by(Contract.updated_at.desc()).limit(8).all()
    total_clients = Client.query.filter_by(created_by=current_user.id).count()
    total_templates = ContractTemplate.query.filter_by(is_active=True, created_by=current_user.id).count()
    recent_activity = AuditLog.query.filter_by(user_id=current_user.id).order_by(AuditLog.created_at.desc()).limit(10).all()

    return render_template('dashboard.html',
        total_contracts=total_contracts,
        by_status=by_status,
        recent_contracts=recent_contracts,
        total_clients=total_clients,
        total_templates=total_templates,
        recent_activity=recent_activity
    )


# ─── Client Routes ────────────────────────────────────────────────────────────

@app.route('/clients')
@login_required
def clients_list():
    q = request.args.get('q', '')
    ctype = request.args.get('type', '')
    query = Client.query.filter_by(created_by=current_user.id)
    if q:
        query = query.filter(
            db.or_(Client.name.ilike(f'%{q}%'), Client.company.ilike(f'%{q}%'), Client.email.ilike(f'%{q}%'))
        )
    if ctype:
        query = query.filter_by(client_type=ctype)
    clients = query.order_by(Client.name).all()
    return render_template('clients/list.html', clients=clients, q=q, ctype=ctype)


@app.route('/clients/new', methods=['GET', 'POST'])
@login_required
def clients_new():
    if request.method == 'POST':
        client = Client(
            name=request.form.get('name', '').strip(),
            company=request.form.get('company', '').strip(),
            email=request.form.get('email', '').strip(),
            phone=request.form.get('phone', '').strip(),
            address=request.form.get('address', '').strip(),
            notes=request.form.get('notes', '').strip(),
            client_type=request.form.get('client_type', 'client'),
            created_by=current_user.id
        )
        if not client.name:
            flash('Name is required.', 'error')
            return render_template('clients/new.html', client=client)

        db.session.add(client)
        log_action('create_client', 'client', None, details=f'Created client: {client.name}')
        db.session.commit()
        flash(f'Client "{client.name}" created successfully.', 'success')
        return redirect(url_for('clients_detail', client_id=client.id))

    return render_template('clients/new.html', client=None)


@app.route('/clients/<int:client_id>')
@login_required
def clients_detail(client_id):
    client = Client.query.get_or_404(client_id)
    if client.created_by != current_user.id:
        abort(403)
    contracts = client.contracts.order_by(Contract.updated_at.desc()).all()
    return render_template('clients/detail.html', client=client, contracts=contracts)


@app.route('/clients/<int:client_id>/edit', methods=['GET', 'POST'])
@login_required
def clients_edit(client_id):
    client = Client.query.get_or_404(client_id)
    if client.created_by != current_user.id:
        abort(403)
    if request.method == 'POST':
        client.name = request.form.get('name', '').strip()
        client.company = request.form.get('company', '').strip()
        client.email = request.form.get('email', '').strip()
        client.phone = request.form.get('phone', '').strip()
        client.address = request.form.get('address', '').strip()
        client.notes = request.form.get('notes', '').strip()
        client.client_type = request.form.get('client_type', 'client')
        client.updated_at = datetime.utcnow()
        log_action('update_client', 'client', client.id, details=f'Updated client: {client.name}')
        db.session.commit()
        flash('Client updated successfully.', 'success')
        return redirect(url_for('clients_detail', client_id=client.id))
    return render_template('clients/new.html', client=client)


# ─── Template Routes ──────────────────────────────────────────────────────────

@app.route('/templates')
@login_required
def templates_list():
    q = request.args.get('q', '')
    ttype = request.args.get('type', '')
    query = ContractTemplate.query.filter_by(is_active=True, created_by=current_user.id)
    if q:
        query = query.filter(ContractTemplate.name.ilike(f'%{q}%'))
    if ttype:
        query = query.filter_by(template_type=ttype)
    templates = query.order_by(ContractTemplate.updated_at.desc()).all()
    return render_template('templates/list.html', templates=templates, q=q, ttype=ttype)


@app.route('/templates/new', methods=['GET', 'POST'])
@login_required
def templates_new():
    if request.method == 'POST':
        content = None

        # File upload takes priority
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'],
                                         f'tpl_{datetime.utcnow().strftime("%Y%m%d%H%M%S")}_{filename}')
                file.save(file_path)
                content = extract_text_from_file(file_path, filename)
            else:
                flash('Unsupported file type. Please upload a .docx or .txt file.', 'error')
                return render_template('templates/new.html', template=None)

        if not content:
            # Raw-HTML editor (used when editing an HTML template)
            raw = request.form.get('content_raw', '').strip()
            content = raw or request.form.get('content', '').strip()

        if not content:
            flash('Template content is required. Enter text or upload a file.', 'error')
            return render_template('templates/new.html', template=None)

        fields = extract_template_fields(content)
        template = ContractTemplate(
            name=request.form.get('name', '').strip(),
            description=request.form.get('description', '').strip(),
            template_type=request.form.get('template_type', 'contract'),
            content=content,
            created_by=current_user.id
        )
        template.fields = fields

        if not template.name:
            flash('Template name is required.', 'error')
            return render_template('templates/new.html', template=template)

        db.session.add(template)
        log_action('create_template', 'template', None, details=f'Created template: {template.name}')
        db.session.commit()
        flash(f'Template "{template.name}" created with {len(fields)} fields detected.', 'success')
        return redirect(url_for('templates_detail', template_id=template.id))

    return render_template('templates/new.html', template=None)


@app.route('/templates/<int:template_id>')
@login_required
def templates_detail(template_id):
    template = ContractTemplate.query.get_or_404(template_id)
    if template.created_by != current_user.id:
        abort(403)
    return render_template('templates/detail.html', template=template)


@app.route('/templates/<int:template_id>/edit', methods=['GET', 'POST'])
@login_required
def templates_edit(template_id):
    template = ContractTemplate.query.get_or_404(template_id)
    if template.created_by != current_user.id:
        abort(403)
    if request.method == 'POST':
        content = None

        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'],
                                         f'tpl_{datetime.utcnow().strftime("%Y%m%d%H%M%S")}_{filename}')
                file.save(file_path)
                content = extract_text_from_file(file_path, filename)
            else:
                flash('Unsupported file type. Please upload a .docx or .txt file.', 'error')
                return render_template('templates/new.html', template=template)

        if not content:
            raw = request.form.get('content_raw', '').strip()
            content = raw or request.form.get('content', '').strip()

        if not content:
            flash('Template content is required.', 'error')
            return render_template('templates/new.html', template=template)

        template.name = request.form.get('name', '').strip()
        template.description = request.form.get('description', '').strip()
        template.template_type = request.form.get('template_type', 'contract')
        template.content = content
        template.fields = extract_template_fields(content)
        template.updated_at = datetime.utcnow()
        log_action('update_template', 'template', template.id, details=f'Updated template: {template.name}')
        db.session.commit()
        flash('Template updated.', 'success')
        return redirect(url_for('templates_detail', template_id=template.id))

    return render_template('templates/new.html', template=template)


@app.route('/templates/<int:template_id>/update-fields', methods=['POST'])
@login_required
def templates_update_fields(template_id):
    """AJAX endpoint — saves the edited HTML content (with field markers) back to the template."""
    template = ContractTemplate.query.get_or_404(template_id)
    if template.created_by != current_user.id:
        abort(403)
    content = request.form.get('content', '').strip()
    if not content:
        return jsonify({'ok': False, 'error': 'Content cannot be empty'}), 400
    # Normalise content before saving so +FIELD+ notation and split markers are
    # always converted to canonical {{FIELD}} form.  This ensures the stored
    # content is consistent with what the browser expects to display.
    if is_html_content(content):
        content = fix_split_placeholders(content)   # handles +FIELD+ and split {{…}}
    else:
        content = _normalize_plus_placeholders(content)
    template.content = content
    template.fields = extract_template_fields(content)
    template.updated_at = datetime.utcnow()
    log_action('update_template', 'template', template.id,
               details=f'Updated fields via field editor: {template.name}')
    db.session.commit()
    return jsonify({'ok': True, 'fields': template.fields})


@app.route('/templates/<int:template_id>/delete', methods=['POST'])
@login_required
def templates_delete(template_id):
    template = ContractTemplate.query.get_or_404(template_id)
    if template.created_by != current_user.id:
        abort(403)
    template.is_active = False
    log_action('delete_template', 'template', template.id, details=f'Deleted template: {template.name}')
    db.session.commit()
    flash('Template deleted.', 'success')
    return redirect(url_for('templates_list'))


@app.route('/api/templates/<int:template_id>/fields')
@login_required
def template_fields_api(template_id):
    template = ContractTemplate.query.get_or_404(template_id)
    if template.created_by != current_user.id:
        abort(403)
    # Always extract live from content so stale fields_json never causes "no fields"
    fields = extract_template_fields(template.content or '')
    # Keep fields_json in sync if they differ
    if fields != template.fields:
        template.fields = fields
        db.session.commit()
    return jsonify({'fields': fields, 'content': template.content})


# ─── Contract Routes ──────────────────────────────────────────────────────────

@app.route('/contracts')
@login_required
def contracts_list():
    q = request.args.get('q', '')
    status = request.args.get('status', '')
    client_id = request.args.get('client_id', '')
    query = Contract.query.filter_by(created_by=current_user.id)
    if q:
        query = query.filter(
            db.or_(Contract.title.ilike(f'%{q}%'), Contract.contract_number.ilike(f'%{q}%'))
        )
    if status:
        query = query.filter_by(status=status)
    if client_id:
        query = query.filter_by(client_id=int(client_id))
    contracts = query.order_by(Contract.updated_at.desc()).all()
    clients = Client.query.filter_by(created_by=current_user.id).order_by(Client.name).all()
    return render_template('contracts/list.html', contracts=contracts, clients=clients,
                           q=q, status=status, client_id=client_id,
                           today=datetime.utcnow().date())


@app.route('/contracts/new', methods=['GET', 'POST'])
@login_required
def contracts_new():
    clients = Client.query.filter_by(created_by=current_user.id).order_by(Client.name).all()
    templates = ContractTemplate.query.filter_by(is_active=True, created_by=current_user.id).order_by(ContractTemplate.name).all()

    if request.method == 'POST':
        template_id = request.form.get('template_id')
        template_id = int(template_id) if template_id else None

        # Validate that the submitted client and template belong to the current user
        submitted_client_id = request.form.get('client_id')
        if not submitted_client_id:
            flash('Client is required.', 'error')
            return render_template('contracts/new.html', clients=clients, templates=templates, contract=None)
        submitted_client = Client.query.filter_by(id=int(submitted_client_id), created_by=current_user.id).first()
        if not submitted_client:
            abort(403)

        if template_id:
            submitted_template = ContractTemplate.query.filter_by(id=template_id, created_by=current_user.id, is_active=True).first()
            if not submitted_template:
                abort(403)

        start_date = None
        end_date = None
        try:
            sd = request.form.get('start_date', '')
            ed = request.form.get('end_date', '')
            if sd:
                start_date = date.fromisoformat(sd)
            if ed:
                end_date = date.fromisoformat(ed)
        except ValueError:
            pass

        value = None
        try:
            v = request.form.get('value', '').strip()
            if v:
                value = float(v.replace(',', ''))
        except ValueError:
            pass

        # Allow uploading already-signed/executed contracts at a specific stage
        _valid_statuses = [
            'draft', 'in_review', 'approved', 'finalized',
            'partially_executed', 'fully_executed', 'expired',
        ]
        initial_status = request.form.get('initial_status', 'draft')
        if initial_status not in _valid_statuses:
            initial_status = 'draft'

        contract = Contract(
            title=request.form.get('title', '').strip(),
            client_id=submitted_client.id,
            template_id=template_id,
            status=initial_status,
            notes=request.form.get('notes', '').strip(),
            start_date=start_date,
            end_date=end_date,
            value=value,
            created_by=current_user.id,
            contract_number=generate_contract_number()
        )
        # Set execution timestamps based on the chosen initial status
        if initial_status in ('finalized', 'partially_executed', 'fully_executed'):
            contract.finalized_at = datetime.utcnow()
        if initial_status == 'fully_executed':
            contract.executed_at = datetime.utcnow()

        if not contract.title or not contract.client_id:
            flash('Title and client are required.', 'error')
            return render_template('contracts/new.html', clients=clients, templates=templates, contract=contract)

        db.session.add(contract)
        # Retry up to 5 times if a concurrent request grabbed the same number
        for _attempt in range(5):
            try:
                db.session.flush()
                break
            except Exception as _flush_err:
                _err = str(_flush_err).lower()
                if 'unique' in _err and 'contract_number' in _err:
                    db.session.rollback()
                    contract.contract_number = generate_contract_number(offset=_attempt + 1)
                    db.session.add(contract)
                else:
                    raise

        # ── Uploaded document (import mode) ──────────────────────────────────
        uploaded_file_name = None
        initial_content = request.form.get('initial_content', '').strip()
        if 'document_file' in request.files and request.files['document_file'].filename:
            file = request.files['document_file']
            if allowed_file(file.filename):
                uploaded_file_name = secure_filename(file.filename)
                ts = datetime.utcnow().strftime('%Y%m%d%H%M%S')
                save_path = os.path.join(
                    app.config['UPLOAD_FOLDER'],
                    f'contract_{contract.id}_{ts}_{uploaded_file_name}'
                )
                file.save(save_path)
                initial_content = extract_text_from_file(save_path, uploaded_file_name)
            else:
                flash('Invalid file type. Allowed: txt, docx, pdf', 'error')
                db.session.rollback()
                return render_template('contracts/new.html', clients=clients, templates=templates, contract=None)

        # ── Template-based content ────────────────────────────────────────────
        if template_id and not initial_content:
            template = ContractTemplate.query.get(template_id)
            if template:
                initial_content = template.content

                # Always extract fields live from the template content so that
                # stale fields_json never causes submitted values to be lost.
                live_fields = extract_template_fields(template.content or '')

                # Save field values for auditing / future reference
                for field_name in live_fields:
                    field_val = request.form.get(f'field_{field_name}', '').strip()
                    if field_val:
                        fv = ContractFieldValue(
                            contract_id=contract.id,
                            field_name=field_name,
                            field_value=field_val
                        )
                        db.session.add(fv)

                # Apply all submitted field_* values to the content, even ones
                # not in live_fields (belt-and-suspenders for edge cases).
                submitted_fields = {
                    k[len('field_'):]: v
                    for k, v in request.form.items()
                    if k.startswith('field_') and v.strip()
                }
                # Merge: live_fields takes priority for ordering, submitted fills values
                field_values = {f: request.form.get(f'field_{f}', '') for f in live_fields}
                field_values.update({k: v for k, v in submitted_fields.items() if k not in field_values})
                initial_content = apply_field_values(initial_content, field_values)

        if initial_content:
            revision = ContractRevision(
                contract_id=contract.id,
                version_number=1,
                content=initial_content,
                changes_summary='Initial version',
                created_by=current_user.id
            )
            db.session.add(revision)

        log_action('create_contract', 'contract', contract.id, contract_id=contract.id,
                   details=f'Created contract: {contract.title}')
        db.session.commit()
        flash(f'Contract "{contract.title}" created.', 'success')
        return redirect(url_for('contracts_detail', contract_id=contract.id))

    return render_template('contracts/new.html', clients=clients, templates=templates, contract=None)


@app.route('/contracts/<int:contract_id>')
@login_required
def contracts_detail(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    revisions = contract.revisions.order_by(ContractRevision.version_number.desc()).all()
    field_values = {fv.field_name: fv.field_value for fv in contract.field_values}
    audit = contract.audit_logs.order_by(AuditLog.created_at.desc()).limit(20).all()
    return render_template('contracts/detail.html', contract=contract, revisions=revisions,
                           field_values=field_values, audit=audit)


@app.route('/contracts/<int:contract_id>/edit', methods=['GET', 'POST'])
@login_required
def contracts_edit(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    clients = Client.query.filter_by(created_by=current_user.id).order_by(Client.name).all()

    if request.method == 'POST':
        # Validate that the submitted client belongs to the current user
        submitted_client_id = request.form.get('client_id')
        if not submitted_client_id:
            flash('Client is required.', 'error')
            return render_template('contracts/edit.html', contract=contract, clients=clients)
        submitted_client = Client.query.filter_by(id=int(submitted_client_id), created_by=current_user.id).first()
        if not submitted_client:
            abort(403)

        contract.title = request.form.get('title', '').strip()
        contract.client_id = submitted_client.id
        contract.status = request.form.get('status', contract.status)
        contract.notes = request.form.get('notes', '').strip()
        contract.updated_at = datetime.utcnow()

        try:
            sd = request.form.get('start_date', '')
            ed = request.form.get('end_date', '')
            contract.start_date = date.fromisoformat(sd) if sd else None
            contract.end_date = date.fromisoformat(ed) if ed else None
        except ValueError:
            pass

        try:
            v = request.form.get('value', '').strip()
            contract.value = float(v.replace(',', '')) if v else None
        except ValueError:
            pass

        if contract.status == 'finalized' and not contract.finalized_at:
            contract.finalized_at = datetime.utcnow()

        log_action('update_contract', 'contract', contract.id, contract_id=contract.id,
                   details=f'Updated contract: {contract.title}, status={contract.status}')
        db.session.commit()
        flash('Contract updated.', 'success')
        return redirect(url_for('contracts_detail', contract_id=contract.id))

    return render_template('contracts/edit.html', contract=contract, clients=clients)


@app.route('/contracts/<int:contract_id>/revision/new', methods=['GET', 'POST'])
@login_required
def revisions_new(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)

    if request.method == 'POST':
        content = request.form.get('content', '').strip()
        changes_summary = request.form.get('changes_summary', '').strip()
        file_name = None

        # Handle file upload
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            if allowed_file(file.filename):
                file_name = secure_filename(file.filename)
                ts = datetime.utcnow().strftime('%Y%m%d%H%M%S')
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], f'rev_{contract_id}_{ts}_{file_name}')
                file.save(file_path)
                content = extract_text_from_file(file_path, file_name)
            else:
                flash('Invalid file type. Allowed: txt, pdf, docx', 'error')
                return redirect(request.url)

        if not content:
            flash('Revision content is required.', 'error')
            return redirect(request.url)

        last_rev = contract.revisions.order_by(ContractRevision.version_number.desc()).first()
        next_version = (last_rev.version_number + 1) if last_rev else 1

        revision = ContractRevision(
            contract_id=contract.id,
            version_number=next_version,
            content=content,
            changes_summary=changes_summary or f'Version {next_version}',
            file_name=file_name,
            created_by=current_user.id
        )
        db.session.add(revision)
        contract.updated_at = datetime.utcnow()

        log_action('add_revision', 'contract_revision', None, contract_id=contract.id,
                   details=f'Added revision v{next_version} to contract {contract.title}')
        db.session.commit()
        flash(f'Revision v{next_version} added.', 'success')
        return redirect(url_for('contracts_detail', contract_id=contract.id))

    latest = contract.latest_revision
    return render_template('contracts/revision_new.html', contract=contract, latest=latest)


@app.route('/contracts/<int:contract_id>/compare')
@login_required
def contracts_compare(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    rev1_id = request.args.get('rev1', type=int)
    rev2_id = request.args.get('rev2', type=int)

    revisions = contract.revisions.order_by(ContractRevision.version_number).all()
    if len(revisions) < 2:
        flash('You need at least 2 revisions to compare.', 'warning')
        return redirect(url_for('contracts_detail', contract_id=contract_id))

    rev1 = ContractRevision.query.get_or_404(rev1_id) if rev1_id else revisions[-2]
    if rev1_id and rev1.contract_id != contract_id:
        abort(403)
    rev2 = ContractRevision.query.get_or_404(rev2_id) if rev2_id else revisions[-1]
    if rev2_id and rev2.contract_id != contract_id:
        abort(403)

    # Generate HTML diff
    differ = difflib.HtmlDiff(wrapcolumn=80)
    diff_html = differ.make_table(
        rev1.content.splitlines(keepends=True),
        rev2.content.splitlines(keepends=True),
        fromdesc=f'Version {rev1.version_number}',
        todesc=f'Version {rev2.version_number}',
        context=True,
        numlines=3
    )

    # Also generate unified diff for stats
    unified = list(difflib.unified_diff(
        rev1.content.splitlines(),
        rev2.content.splitlines()
    ))
    added = sum(1 for l in unified if l.startswith('+') and not l.startswith('+++'))
    removed = sum(1 for l in unified if l.startswith('-') and not l.startswith('---'))

    return render_template('contracts/compare.html', contract=contract, rev1=rev1, rev2=rev2,
                           revisions=revisions, diff_html=diff_html, added=added, removed=removed)


@app.route('/contracts/<int:contract_id>/finalize', methods=['POST'])
@login_required
def contracts_finalize(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    revision_id = request.form.get('revision_id', type=int)

    if revision_id:
        revision = ContractRevision.query.get_or_404(revision_id)
        if revision.contract_id != contract_id:
            abort(403)
    else:
        revision = contract.latest_revision

    if not revision:
        flash('No revision found to finalize.', 'error')
        return redirect(url_for('contracts_detail', contract_id=contract_id))

    # Mark revision as finalized
    revision.is_finalized = True
    contract.status = 'finalized'
    contract.finalized_at = datetime.utcnow()
    contract.updated_at = datetime.utcnow()

    log_action('finalize_contract', 'contract', contract.id, contract_id=contract.id,
               details=f'Finalized contract: {contract.title}, revision v{revision.version_number}')
    db.session.commit()
    flash(f'Contract "{contract.title}" has been finalized (v{revision.version_number}).', 'success')
    return redirect(url_for('contracts_detail', contract_id=contract_id))


@app.route('/contracts/<int:contract_id>/pdf')
@login_required
def contracts_pdf(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    revision_id = request.args.get('revision_id', type=int)
    if revision_id:
        revision = ContractRevision.query.get_or_404(revision_id)
        if revision.contract_id != contract_id:
            abort(403)
    else:
        revision = contract.latest_revision

    pdf_buffer = generate_pdf(contract, revision)
    filename = f'{contract.contract_number}_{contract.title.replace(" ", "_")}'
    if revision:
        filename += f'_v{revision.version_number}'
    filename += '.pdf'

    log_action('export_pdf', 'contract', contract.id, contract_id=contract.id,
               details=f'Exported PDF for contract: {contract.title}')
    db.session.commit()

    return send_file(pdf_buffer, as_attachment=True, download_name=filename,
                     mimetype='application/pdf')


@app.route('/contracts/<int:contract_id>/docx')
@login_required
def contracts_docx(contract_id):
    contract = Contract.query.filter_by(id=contract_id, created_by=current_user.id).first_or_404()
    revision_id = request.args.get('revision_id', type=int)
    if revision_id:
        revision = ContractRevision.query.get_or_404(revision_id)
        if revision.contract_id != contract_id:
            abort(403)
    else:
        revision = contract.latest_revision

    if not DOCX_AVAILABLE:
        flash('Word document export is not available (python-docx not installed).', 'error')
        return redirect(url_for('contracts_detail', contract_id=contract_id))

    docx_buffer = generate_docx(contract, revision)
    if not docx_buffer:
        flash('Could not generate Word document.', 'error')
        return redirect(url_for('contracts_detail', contract_id=contract_id))

    safe_title = re.sub(r'[^\w\s-]', '', contract.title).strip().replace(' ', '_')
    filename = f'{contract.contract_number}_{safe_title}'
    if revision:
        filename += f'_v{revision.version_number}'
    filename += '.docx'

    log_action('export_docx', 'contract', contract.id, contract_id=contract.id,
               details=f'Exported DOCX for contract: {contract.title}')
    db.session.commit()

    return send_file(
        docx_buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/contracts/<int:contract_id>/status', methods=['POST'])
@login_required
def contracts_status(contract_id):
    contract = Contract.query.get_or_404(contract_id)
    if contract.created_by != current_user.id:
        abort(403)
    new_status = request.form.get('status')
    valid_statuses = [
        'draft', 'in_review', 'approved', 'finalized',
        'partially_executed', 'fully_executed', 'expired',
    ]
    if new_status in valid_statuses:
        old_status = contract.status
        contract.status = new_status
        contract.updated_at = datetime.utcnow()
        # Set finalized_at the first time the contract enters the signature pipeline
        if new_status in ('finalized', 'partially_executed', 'fully_executed'):
            if not contract.finalized_at:
                contract.finalized_at = datetime.utcnow()
        # Record the moment of full execution
        if new_status == 'fully_executed' and not contract.executed_at:
            contract.executed_at = datetime.utcnow()
        log_action('status_change', 'contract', contract.id, contract_id=contract.id,
                   details=f'Status changed: {old_status} → {new_status}')
        db.session.commit()
        labels = {
            'draft': 'Draft', 'in_review': 'In Review', 'approved': 'Approved',
            'finalized': 'Finalized — Sent for Signature',
            'partially_executed': 'Partially Executed',
            'fully_executed': 'Fully Executed',
            'expired': 'Expired',
        }
        flash(f'Status updated to {labels.get(new_status, new_status)}.', 'success')
    return redirect(url_for('contracts_detail', contract_id=contract_id))


@app.route('/contracts/<int:contract_id>/delete', methods=['POST'])
@login_required
def contracts_delete(contract_id):
    contract = Contract.query.filter_by(id=contract_id, created_by=current_user.id).first_or_404()
    title = contract.title

    # Remove uploaded revision files from disk
    for rev in contract.revisions.all():
        if rev.file_path and os.path.exists(rev.file_path):
            try:
                os.remove(rev.file_path)
            except OSError:
                pass

    # Delete child rows before deleting the contract (FK constraints)
    ContractFieldValue.query.filter_by(contract_id=contract.id).delete()
    ContractRevision.query.filter_by(contract_id=contract.id).delete()
    # Null-out audit log references so the deletion history is preserved
    AuditLog.query.filter_by(contract_id=contract.id).update({'contract_id': None})

    log_action('delete_contract', 'contract', contract.id,
               details=f'Deleted contract: {title} ({contract.contract_number})')
    db.session.delete(contract)
    db.session.commit()
    flash(f'Contract "{title}" has been permanently deleted.', 'success')
    return redirect(url_for('contracts_list'))


# ─── Contract Report Export ───────────────────────────────────────────────────

@app.route('/contracts/export')
@login_required
def contracts_export():
    """Download all contracts as a CSV report with values and monthly revenue."""
    import csv
    from io import StringIO

    contracts = (
        Contract.query
        .filter_by(created_by=current_user.id)
        .order_by(Contract.status, Contract.contract_number)
        .all()
    )

    si = StringIO()
    writer = csv.writer(si)
    writer.writerow([
        'Contract #', 'Title', 'Client', 'Company', 'Status',
        'Start Date', 'End Date', 'Term (months)',
        'Total Value ($)', 'Monthly Revenue ($)',
        'Finalized Date', 'Executed Date',
    ])

    for c in contracts:
        # Term length in whole months
        months = 0
        if c.start_date and c.end_date:
            months = max(0,
                (c.end_date.year  - c.start_date.year)  * 12 +
                (c.end_date.month - c.start_date.month)
            )

        total_val   = float(c.value) if c.value else None
        monthly_rev = (total_val / months) if (total_val and months) else None

        writer.writerow([
            c.contract_number,
            c.title,
            c.client.name,
            c.client.company or '',
            c.status.replace('_', ' ').title(),
            c.start_date.strftime('%Y-%m-%d')  if c.start_date  else '',
            c.end_date.strftime('%Y-%m-%d')    if c.end_date    else '',
            months or '',
            f'{total_val:.2f}'   if total_val   is not None else '',
            f'{monthly_rev:.2f}' if monthly_rev is not None else '',
            c.finalized_at.strftime('%Y-%m-%d') if c.finalized_at else '',
            c.executed_at.strftime('%Y-%m-%d')  if c.executed_at  else '',
        ])

    log_action('export_report', 'contract', details='Exported contracts CSV report')

    output = si.getvalue()
    response = make_response(output)
    response.headers['Content-Type'] = 'text/csv; charset=utf-8'
    ts = datetime.utcnow().strftime('%Y%m%d')
    response.headers['Content-Disposition'] = (
        f'attachment; filename=contracts_report_{ts}.csv'
    )
    return response


# ─── Template Parsing API ─────────────────────────────────────────────────────

@app.route('/api/detect-fields', methods=['POST'])
@login_required
def detect_fields_api():
    content = request.json.get('content', '')
    fields = extract_template_fields(content)
    return jsonify({'fields': fields})


# ─── Init ─────────────────────────────────────────────────────────────────────

# Create tables on startup regardless of how the app is launched (gunicorn or
# direct). This is idempotent — SQLAlchemy skips tables that already exist.
with app.app_context():
    db.create_all()
    # Add new columns introduced after initial deployment (safe no-op if they already exist)
    try:
        db.session.execute(db.text(
            'ALTER TABLE contracts ADD COLUMN IF NOT EXISTS executed_at TIMESTAMP'
        ))
        db.session.commit()
    except Exception:
        db.session.rollback()


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=False)
